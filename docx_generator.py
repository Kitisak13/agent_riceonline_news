# docx_generator.py - Document Generation Module for Rice News Aggregator
# =========================================================================

from typing import Any, Dict, List
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from utils import logger


def set_font_style(
    run: Any,
    font_name: str = 'Times New Roman',
    size: int = 11,
    bold: bool = False,
    italic: bool = False
) -> None:
    """
    Set explicit font formatting on a docx Run element.
    """
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    fonts.set(qn('w:eastAsia'), font_name)
    fonts.set(qn('w:cs'), font_name)
    rPr.append(fonts)


def create_document(news_list: List[Dict[str, Any]], filename: str) -> str:
    """
    Generate a formatted Word document (.docx) from processed news items.
    """
    logger.info(f"📄 Generating document: {filename}")
    doc = Document()
    success_count = 0
    failed_count = 0
    
    for news in news_list:
        headline = news.get('headline', 'Untitled')
        url = news.get('url') or news.get('URL', '')
        
        if news.get('processed_data'):
            data = news['processed_data']
            
            p_headline = doc.add_paragraph()
            run_headline = p_headline.add_run(data.get('cleaned_headline', headline))
            set_font_style(run_headline, size=12, bold=True)
            
            p_date = doc.add_paragraph()
            run_date = p_date.add_run(f"Date: {data.get('date_str', 'N/A')}")
            set_font_style(run_date, size=10, italic=True)
            
            content = data.get('full_content', '')
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    p_content = doc.add_paragraph()
                    p_content.paragraph_format.space_after = Pt(6)
                    run_content = p_content.add_run(para)
                    set_font_style(run_content, size=11)
            
            if url:
                p_link = doc.add_paragraph()
                run_link = p_link.add_run(url)
                set_font_style(run_link, size=11)
            
            success_count += 1
        else:
            p_headline = doc.add_paragraph()
            run_headline = p_headline.add_run(headline)
            set_font_style(run_headline, size=12, bold=True)
            
            p_fail = doc.add_paragraph()
            run_fail = p_fail.add_run("[Cannot find content]")
            set_font_style(run_fail, size=11, italic=True)
            
            if url:
                p_link = doc.add_paragraph()
                run_link = p_link.add_run(url)
                set_font_style(run_link, size=11)
                
            failed_count += 1
            
        doc.add_paragraph()
        
    doc.save(filename)
    logger.info(f"✅ Saved to {filename}: {success_count} success, {failed_count} failed")
    return filename
