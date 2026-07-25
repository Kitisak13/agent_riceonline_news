# main.py - Report Generation Module (Editor & Production) - Free AI V1
# ======================================================================

import json
import os
import random
import smtplib
import time
import atexit
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from urllib.parse import quote

import requests
import trafilatura
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google import genai
from google.genai import types

# Try importing curl_cffi for Cloudflare bypass, fallback to None
try:
    from curl_cffi import requests as curl_requests
    logger_curl = True
except ImportError:
    curl_requests = None
    logger_curl = False

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait

from config import (
    HEADLESS_BROWSER,
    MAX_RETRIES,
    GEMINI_PRIMARY_MODEL,
    GEMINI_FALLBACK_MODEL,
    OLD_NEWS_DAYS,
    PAGE_LOAD_TIMEOUT,
    REQUIRED_ENV_KEYS,
    SELENIUM_SESSION_MAX_URLS,
    SELENIUM_WAIT_TIMEOUT,
    USER_AGENT,
)
from utils import (
    DomainFailureCache,
    format_duration,
    is_old_news,
    is_valid_url,
    logger,
    mask_sensitive_value,
    parse_date_flexible,
    requires_selenium,
    resolve_google_redirect,
    gemini_limiter,
)

# Load environment variables
load_dotenv()
if not os.getenv("GEMINI_API_KEY"):
    load_dotenv("../.env")

INPUT_FILE = 'source.json'
TEST_MODE = False

# User agent rotation
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/133.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/132.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/133.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:134.0) Gecko/20100101 Firefox/134.0",
]

def get_random_user_agent():
    return random.choice(USER_AGENTS)

def get_google_cache_url(url):
    return f"https://webcache.googleusercontent.com/search?q=cache:{quote(url)}"

# Configure Gemini Client
gemini_client = None
if os.getenv("GEMINI_API_KEY"):
    try:
        clean_key = os.getenv("GEMINI_API_KEY").strip().strip('"').strip("'")
        gemini_client = genai.Client(api_key=clean_key)
        logger.info(f"✨ Gemini Client configured for editing and fallbacks.")
    except Exception as e:
        logger.error(f"Failed to configure Gemini: {e}")


# ==============================================================================
# SELENIUM SESSION MANAGER
# ==============================================================================
class SeleniumSessionManager:
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._driver = None
            cls._instance._url_count = 0
            cls._instance._max_urls = SELENIUM_SESSION_MAX_URLS
        return cls._instance
    
    def _create_driver(self):
        def get_chrome_path():
            paths = [
                r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
                r"C:\Users\%USERNAME%\AppData\Local\Google\Chrome\Application\chrome.exe",
            ]
            for path in paths:
                expanded = os.path.expandvars(path)
                if os.path.exists(expanded):
                    return expanded
            return None

        chrome_binary = get_chrome_path()
        chrome_options = Options()
        if chrome_binary:
            chrome_options.binary_location = chrome_binary
            
        if HEADLESS_BROWSER:
            chrome_options.add_argument("--headless=new")
        
        chrome_options.add_argument("--disable-gpu")
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--mute-audio")
        chrome_options.add_argument(f"--user-agent={get_random_user_agent()}")
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)
        
        try:
            logger.info("   📥 Auto-detecting ChromeDriver version via webdriver-manager...")
            service = Service(ChromeDriverManager().install())
        except Exception as e:
            logger.warning(f"   ⚠️ webdriver-manager failed ({e}), trying local chromedriver...")
            local_driver_path = os.path.join(os.getcwd(), "chromedriver.exe")
            if os.path.exists(local_driver_path):
                service = Service(executable_path=local_driver_path)
            else:
                # Search parent directory as well
                parent_driver = os.path.join(os.getcwd(), "..", "chromedriver.exe")
                if os.path.exists(parent_driver):
                    service = Service(executable_path=parent_driver)
                else:
                    raise RuntimeError("No compatible ChromeDriver found. Please update chromedriver.exe.")
            
        driver = webdriver.Chrome(service=service, options=chrome_options)
        driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
            "source": "Object.defineProperty(navigator, 'webdriver', { get: () => undefined })"
        })
        driver.set_page_load_timeout(PAGE_LOAD_TIMEOUT)
        driver.set_script_timeout(PAGE_LOAD_TIMEOUT)
        return driver
    
    def _get_driver(self):
        if self._url_count >= self._max_urls and self._driver:
            logger.info(f"   🔄 Recycling browser session (after {self._url_count} URLs)")
            self._close_driver()
        
        if self._driver is None:
            logger.info("   🌐 Creating new browser session...")
            self._driver = self._create_driver()
            self._url_count = 0
        
        return self._driver
    
    def _close_driver(self):
        if self._driver:
            try:
                self._driver.quit()
            except Exception:
                pass
            self._driver = None
            self._url_count = 0
    
    def _wait_for_content(self, driver):
        try:
            WebDriverWait(driver, SELENIUM_WAIT_TIMEOUT).until(
                lambda d: len(d.find_element(By.TAG_NAME, "body").text) > 200
            )
            return True
        except Exception:
            return False
    
    def _is_cloudflare_blocked(self, driver):
        try:
            page_text = driver.find_element(By.TAG_NAME, "body").text.lower()
            blocked_indicators = [
                "checking your browser", "please wait", "cloudflare", "just a moment",
                "ray id", "ddos protection", "access denied", "403 forbidden"
            ]
            return any(ind in page_text for ind in blocked_indicators)
        except Exception:
            return False

    def _expand_hidden_content(self, driver):
        try:
            keywords = ["read more", "continue reading", "show more", "load more", "expand", "view full article"]
            xpath_parts = [
                f"contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{kw}')" 
                for kw in keywords
            ]
            xpath_query = f"//button[{' or '.join(xpath_parts)}] | //a[{' or '.join(xpath_parts)}]"
            buttons = driver.find_elements(By.XPATH, xpath_query)
            
            for btn in buttons:
                if not btn.is_displayed(): continue
                text = btn.text.strip().lower()
                if len(text) > 30: continue
                try:
                    driver.execute_script("arguments[0].scrollIntoView(true);", btn)
                    time.sleep(0.5)
                    driver.execute_script("arguments[0].click();", btn)
                    logger.info(f"   🖱️ Auto-clicked expand button: '{text}'")
                    time.sleep(1.0)
                    break
                except Exception:
                    continue
        except Exception:
            pass

    def scrape_url(self, url, use_cache_fallback=True):
        try:
            driver = self._get_driver()
            logger.info(f"   🌐 Selenium [{self._url_count + 1}/{self._max_urls}]: {url[:50]}...")
            time.sleep(random.uniform(0.5, 1.5))
            
            driver.get(url)
            self._url_count += 1
            
            content_loaded = self._wait_for_content(driver)
            
            if self._is_cloudflare_blocked(driver):
                logger.warning("   🛡️ Cloudflare detected in Selenium!")
                if use_cache_fallback:
                    logger.info("   📦 Trying Google Cache fallback...")
                    cache_url = get_google_cache_url(url)
                    time.sleep(random.uniform(1, 2))
                    driver.get(cache_url)
                    self._url_count += 1
                    if not self._wait_for_content(driver):
                        logger.error("   ❌ Cache fallback failed")
                        return None
                else:
                    return None
            
            self._expand_hidden_content(driver)
            html = driver.page_source
            
            # Try Trafilatura on HTML
            extract_params = {'include_comments': False, 'include_tables': True, 'output_format': 'json', 'with_metadata': True}
            extract_str = trafilatura.extract(html, **extract_params)
            if extract_str:
                return json.loads(extract_str).get('text', '')
            
            # Fallback to body.text
            body = driver.find_element(By.TAG_NAME, "body")
            text = body.text.strip()
            return text if len(text) > 200 else None
            
        except Exception as e:
            error_msg = str(e).lower()
            if any(term in error_msg for term in ["timeout", "max retries", "connectionpool", "not reachable"]):
                logger.error(f"   ❌ Critical Driver Error: {str(e)[:100]}")
                self._close_driver()
            else:
                logger.error(f"   ❌ Selenium Error: {e}")
                self._close_driver()
            return None
    
    def close(self):
        self._close_driver()
        logger.info("   🔒 Browser session closed")


_selenium_manager = None
def get_selenium_manager():
    global _selenium_manager
    if _selenium_manager is None:
        _selenium_manager = SeleniumSessionManager()
    return _selenium_manager

def cleanup_selenium():
    global _selenium_manager
    if _selenium_manager:
        _selenium_manager.close()
        _selenium_manager = None

atexit.register(cleanup_selenium)


# ==============================================================================
# PIPELINE STAGES
# ==============================================================================

def check_system_health():
    logger.info("=" * 60)
    logger.info("🕵️ SYSTEM DIAGNOSTIC (Free AI V1)")
    logger.info(f"   Working Directory: {os.getcwd()}")
    logger.info(f"   curl_cffi integration: {'✅ Loaded' if logger_curl else '❌ Not available (falls back to requests)'}")
    
    missing_keys = []
    for key in REQUIRED_ENV_KEYS:
        value = os.getenv(key)
        if value:
            logger.info(f"   ✅ {key}: {mask_sensitive_value(value)}")
        else:
            logger.error(f"   ❌ {key}: NOT FOUND")
            missing_keys.append(key)
            
    cache = DomainFailureCache()
    logger.info(f"   📦 Domain Cache: {cache.get_stats()['selenium_required_count']} domains require Selenium")
    logger.info("=" * 60)
    return len(missing_keys) == 0


def gemini_grounding_fallback_scrape(url):
    """
    Fallback Layer 5: If all HTTP & Selenium scrapers fail (due to Cloudflare, bot protections, etc.),
    use Gemini 2.5 Flash Search Grounding (20 RPD) to search Google for the exact URL and extract text.
    """
    if not gemini_client:
        return None
        
    logger.info(f"   🧠 Fallback Layer 5: Attempting Gemini Grounding Scraper for {url[:50]}...")
    
    prompt = f"""
    Search Google for this exact URL: "{url}"
    Read the content of that article from the search index / web cache.
    Extract the COMPLETE text of the news article.
    
    Rules:
    1. DO NOT SUMMARIZE. Extract the full article.
    2. DO NOT REPHRASE. Maintain the original text.
    3. Remove ads, navigation elements, footers, headers.
    4. Respond directly with the article text. If you cannot find the content, output "None".
    """
    
    try:
        # Respect rate limits
        gemini_limiter.wait()
        
        config = types.GenerateContentConfig(
            tools=[types.Tool(google_search=types.GoogleSearch())]
        )
        
        # We must use gemini-2.5-flash since 3.1-flash-lite doesn't have Search Grounding quota (0 RPD)
        response = gemini_client.models.generate_content(
            model=GEMINI_FALLBACK_MODEL,
            contents=prompt,
            config=config
        )
        
        result_text = response.text
        if not result_text:
            logger.warning("   ❌ Gemini Grounding Scraper returned no content or None.")
            return None
        result_text = result_text.strip()
        if "None" in result_text or len(result_text) < 200:
            logger.warning("   ❌ Gemini Grounding Scraper returned no content or None.")
            return None
            
        logger.info(f"   ✅ Gemini Grounding Scraper Success! ({len(result_text)} chars)")
        return {"text": result_text, "title": "", "date": "", "source": url}
        
    except Exception as e:
        logger.error(f"   ⚠️ Gemini Grounding Scraper Failed: {e}")
        return None


def scrape_content(url):
    """
    Multi-layer Scraping Pipeline.
    """
    domain_cache = DomainFailureCache()
    
    # Layer 0: Check fast-track list
    if requires_selenium(url):
        logger.info(f"   ⚡ Fast-track to Selenium (known domain)")
        manager = get_selenium_manager()
        text = manager.scrape_url(url, use_cache_fallback=True)
        if text:
            return {"text": text, "title": "", "date": "", "source": url}
        return None
        
    extract_params = {'include_comments': False, 'include_tables': True, 'output_format': 'json', 'with_metadata': True}
    
    # Layer 1: Try curl_cffi (if available) -> bypasses Cloudflare extremely well
    if curl_requests:
        try:
            logger.debug(f"   尝试 curl_cffi: {url[:50]}...")
            headers = {
                'User-Agent': get_random_user_agent(),
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            }
            r = curl_requests.get(url, headers=headers, impersonate="chrome110", timeout=15)
            if r.status_code == 200:
                extract_str = trafilatura.extract(r.text, **extract_params)
                if extract_str:
                    logger.info("   ✅ curl_cffi + Trafilatura success!")
                    domain_cache.record_success(url)
                    return json.loads(extract_str)
        except Exception as e:
            logger.debug(f"   curl_cffi failed: {e}")

    # Layer 2: Trafilatura direct fetch
    try:
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            extract_str = trafilatura.extract(downloaded, **extract_params)
            if extract_str:
                logger.info("   ✅ Trafilatura direct success!")
                domain_cache.record_success(url)
                return json.loads(extract_str)
    except Exception as e:
        logger.debug(f"   Trafilatura direct failed: {e}")

    # Layer 3: Requests + BeautifulSoup Fallback
    try:
        headers = {'User-Agent': get_random_user_agent()}
        response = requests.get(url, headers=headers, timeout=15)
        if response.status_code == 200:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.text, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'iframe', 'noscript']):
                tag.decompose()
                
            article = (
                soup.find('article') or 
                soup.find('div', class_=lambda c: c and any(k in str(c).lower() for k in ['article', 'content', 'story', 'post-body'])) or
                soup.find('main')
            )
            target = article if article else soup.body
            if target:
                paragraphs = target.find_all('p')
                text_parts = [p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 30]
                if len(text_parts) >= 3:
                    logger.info("   ✅ BeautifulSoup fallback success!")
                    domain_cache.record_success(url)
                    return {
                        'text': '\n\n'.join(text_parts),
                        'title': soup.title.string if soup.title else '',
                        'date': '',
                        'source': url,
                    }
    except Exception:
        pass

    # Layer 4: Selenium
    logger.info("   ⚠️ HTTP extraction failed, trying Selenium...")
    domain_cache.record_failure(url)
    manager = get_selenium_manager()
    sel_text = manager.scrape_url(url, use_cache_fallback=True)
    if sel_text:
        return {"text": sel_text, "title": "", "date": "", "source": url}

    # Layer 5: Gemini Search Grounding Fallback
    fallback_data = gemini_grounding_fallback_scrape(url)
    if fallback_data:
        return fallback_data

    return None


def process_with_ai(headline, content, metadata_date=None):
    """
    Clean and format content using gemini-3.1-flash-lite (completely free, 500 RPD).
    """
    if not gemini_client:
        return None
        
    content_safe = content[:12000]
    date_context = f"Metadata Date: {metadata_date}" if metadata_date else "No metadata date."
    
    prompt = f"""Role: Senior News Editor.
    
    TASK: Clean and format the raw article text below.
    RULES:
    1. DO NOT SUMMARIZE. You must retain every paragraph and factual details.
    2. DO NOT REPHRASE. Keep original wording.
    3. REMOVE ONLY: Ads, social media widgets, navigation/headers/footers, and links to unrelated stories.
    4. FORMAT: Use double newlines (\\n\\n) between paragraphs.
    5. If the raw content does not contain the actual news article corresponding to the headline (e.g. it is just page navigation, a list of other news, a cookie consent warning, or a video player with no text transcript), set "full_content" to "(Can not find the content)".
    6. Output must be valid JSON matching the format below. Do not add markdown backticks wrapper around the JSON.
    
    Headline: {headline}
    {date_context}
    
    Raw Content:
    {content_safe}
    
    Respond strictly in this JSON format:
    {{
      "cleaned_headline": "Original headline",
      "full_content": "Cleaned article text here...",
      "date_str": "YYYY-MM-DD or null"
    }}"""

    for attempt in range(MAX_RETRIES):
        try:
            logger.debug(f"   🤖 AI Call (attempt {attempt+1}/3)...")
            
            # Rate limiting
            gemini_limiter.wait()
            
            response = gemini_client.models.generate_content(
                model=GEMINI_PRIMARY_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.1,
                )
            )
            
            raw_response = response.text
            if not raw_response:
                raise ValueError("Response text is empty or None")
            raw_response = raw_response.strip()
            
            try:
                data = json.loads(raw_response)
                full_content = data.get("full_content", "")
                if not full_content:
                    raise ValueError("Content too short or empty")
                # If the AI explicitly returned "(Can not find the content)" or "None", it is a valid finding, not a parse error
                if full_content in ["(Can not find the content)", "None"]:
                    return data
                if len(full_content) < 100:
                    raise ValueError("Content too short")
                return data
            except (json.JSONDecodeError, ValueError) as je:
                logger.warning(f"   ⚠️ AI attempt {attempt+1} JSON parse error: {je}")
                
        except Exception as e:
            logger.warning(f"   ⚠️ AI attempt {attempt+1} failed: {e}")
            
        time.sleep(2)
        
    return None


def process_single_item(item):
    headline = item.get('headline', 'Unknown')
    url = item.get('URL')
    
    logger.info(f"🚀 Processing: {headline[:50]}...")
    
    if not is_valid_url(url):
        logger.info(f"   🚫 Skipped (invalid/blocked URL)")
        item['processed_data'] = None
        item['fail_reason'] = "Invalid URL / Blocked Domain"
        return item

    # Try normal scraping first
    scraped_data = scrape_content(url)
    item['url'] = url
    
    text_content = ""
    metadata_date = None
    if scraped_data:
        text_content = scraped_data.get('text', '') if isinstance(scraped_data, dict) else scraped_data
        metadata_date = scraped_data.get('date') if isinstance(scraped_data, dict) else None
    
    # Check if we got a valid scrape (not empty, not too short, not a bot block message)
    is_valid_scrape = False
    if text_content and len(text_content.strip()) > 300:
        lower_content = text_content.lower()
        block_indicators = ["javascript must be enabled", "enable javascript", "access denied", "please wait", "cloudflare", "attention required"]
        if not any(ind in lower_content for ind in block_indicators):
            is_valid_scrape = True
            
    ai_result = None
    if is_valid_scrape:
        if metadata_date and is_old_news(metadata_date, days=OLD_NEWS_DAYS):
            logger.info(f"   ⏳ Old news ({metadata_date}) - skipped")
            item['processed_data'] = {
                "cleaned_headline": headline,
                "full_content": "(Can not find the content)",
                "date_str": metadata_date
            }
            return item
            
        ai_result = process_with_ai(headline, text_content, metadata_date)

    # Fallback to Gemini Grounding Scraper if normal scrape failed, was blocked, or AI returned no content
    if (not ai_result or 
        not ai_result.get('full_content') or 
        len(ai_result['full_content'].strip()) < 100 or 
        ai_result['full_content'] == "(Can not find the content)"):
        
        logger.warning("   ⚠️ Scraped text is empty, blocked, or AI cleaning failed. Trying Gemini Grounding fallback...")
        fallback_data = gemini_grounding_fallback_scrape(url)
        if fallback_data:
            fallback_text = fallback_data.get('text', '')
            fallback_date = fallback_data.get('date')
            if fallback_text and len(fallback_text.strip()) > 100:
                ai_result = process_with_ai(headline, fallback_text, fallback_date)

    if ai_result:
        if not ai_result.get('date_str') and metadata_date:
            ai_result['date_str'] = metadata_date
        if is_old_news(ai_result.get('date_str'), days=OLD_NEWS_DAYS):
            logger.info(f"   ⏳ Old news detected post-AI")
            ai_result['full_content'] = "(Can not find the content)"
        item['processed_data'] = ai_result
    else:
        logger.warning(f"   ❌ All scraping and AI fallback options failed")
        item['processed_data'] = None
        item['fail_reason'] = "Scraping/AI failed"
        
    return item


def set_font_style(run, font_name='Times New Roman', size=11, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    
    r = run._element
    rPr = r.get_or_add_rPr()
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), font_name)
    fonts.set(qn('w:hAnsi'), font_name)
    fonts.set(qn('w:eastAsia'), font_name)
    fonts.set(qn('w:cs'), font_name)
    rPr.append(fonts)


def create_document(news_list, filename):
    logger.info(f"📄 Generating document: {filename}")
    doc = Document()
    success_count = 0
    failed_count = 0
    
    for news in news_list:
        headline = news.get('headline', 'Untitled')
        url = news.get('url') or news.get('URL', '')
        
        if news.get('processed_data'):
            data = news['processed_data']
            
            p_headline = doc.add_paragraph()
            run_headline = p_headline.add_run(data.get('cleaned_headline', headline))
            set_font_style(run_headline, size=12, bold=True)
            
            p_date = doc.add_paragraph()
            run_date = p_date.add_run(f"Date: {data.get('date_str', 'N/A')}")
            set_font_style(run_date, size=10, italic=True)
            
            content = data.get('full_content', '')
            for para in content.split('\n'):
                para = para.strip()
                if para:
                    p_content = doc.add_paragraph()
                    p_content.paragraph_format.space_after = Pt(6)
                    run_content = p_content.add_run(para)
                    set_font_style(run_content, size=11)
            
            if url:
                p_link = doc.add_paragraph()
                run_link = p_link.add_run(url)
                set_font_style(run_link, size=11)
            
            success_count += 1
        else:
            p_headline = doc.add_paragraph()
            run_headline = p_headline.add_run(headline)
            set_font_style(run_headline, size=12, bold=True)
            
            p_fail = doc.add_paragraph()
            run_fail = p_fail.add_run("[Cannot find content]")
            set_font_style(run_fail, size=11, italic=True)
            
            if url:
                p_link = doc.add_paragraph()
                run_link = p_link.add_run(url)
                set_font_style(run_link, size=11)
                
            failed_count += 1
            
        doc.add_paragraph()
        
    doc.save(filename)
    logger.info(f"✅ Saved to {filename}: {success_count} success, {failed_count} failed")
    return filename


def upload_to_drive(filename):
    folder_id = os.getenv("DRIVE_FOLDER_ID")
    creds_json = os.getenv("GOOGLE_CREDENTIALS_JSON")
    
    if not folder_id or not creds_json:
        logger.warning("Skipping Drive upload: Credentials missing")
        return None

    logger.info(f"☁️ Uploading to Google Drive: {filename}")
    try:
        creds_dict = json.loads(creds_json)
        creds = service_account.Credentials.from_service_account_info(
            creds_dict, scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=creds)
        
        file_metadata = {'name': filename, 'parents': [folder_id]}
        media = MediaFileUpload(
            filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        file_id = file.get('id')
        logger.info(f"✅ Drive Upload Success! ID={file_id}")
        return file_id
    except Exception as e:
        logger.error(f"Drive upload failed: {e}")
        return None


def send_email(filename, recipient):
    sender_email = os.getenv("EMAIL_SENDER")
    sender_password = os.getenv("EMAIL_PASSWORD")
    
    if not sender_email or not sender_password:
        logger.warning("Skipping email: Credentials missing")
        return False

    logger.info(f"📧 Sending email report to: {recipient}")
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient
        msg['Subject'] = f"Rice News Report: {datetime.now().strftime('%Y-%m-%d')}"
        
        body = "Attached is the Daily Rice News Report generated by the automated script."
        msg.attach(MIMEText(body, 'plain'))

        with open(filename, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(part)

        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient, msg.as_string())
            
        logger.info("✅ Email sent successfully.")
        return True
    except Exception as e:
        logger.error(f"Email delivery failed: {e}")
        return False


def main():
    start_time = time.time()
    
    logger.info("=" * 60)
    logger.info("📰 RICE NEWS REPORT GENERATOR V1 - Starting")
    logger.info("=" * 60)
    
    if not check_system_health():
        logger.warning("Some environment variables are missing. Proceeding with caution.")
        
    if not gemini_client:
        logger.critical("AI Client failed to load. Aborting.")
        return

    if not os.path.exists(INPUT_FILE):
        logger.critical(f"Input file not found: {INPUT_FILE}")
        return

    try:
        with open(INPUT_FILE, 'r', encoding='utf-8') as f:
            headlines = json.load(f)
        logger.info(f"📂 Loaded {len(headlines)} items from {INPUT_FILE}")
    except Exception as e:
        logger.critical(f"Failed to read {INPUT_FILE}: {e}")
        return

    target_headlines = headlines[-7:] if TEST_MODE else headlines
    logger.info(f"🎯 Processing {len(target_headlines)} items sequentially for stability...")

    results = []
    failed_items = []
    
    for idx, item in enumerate(target_headlines, 1):
        try:
            logger.info(f"\n--- Item [{idx}/{len(target_headlines)}] ---")
            res = process_single_item(item)
            results.append(res)
            
            if not res.get('processed_data'):
                reason = res.get('fail_reason', 'Scraping/AI failure')
                failed_items.append({"headline": item.get('headline'), "reason": reason})
        except Exception as e:
            logger.error(f"Error processing item {idx}: {e}")
            failed_items.append({"headline": item.get('headline'), "reason": f"Fatal Error: {e}"})

    # Cleanup Selenium session
    cleanup_selenium()

    # Generate document and outputs
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_filename = f"RiceNews_Report_{timestamp}.docx"
    
    if results:
        create_document(results, output_filename)
        upload_to_drive(output_filename)
        target_email = os.getenv("EMAIL_SENDER")
        if target_email:
            send_email(output_filename, target_email)
    else:
        logger.warning("No articles processed successfully.")

    duration = format_duration(time.time() - start_time)
    logger.info("\n" + "=" * 60)
    logger.info("🎉 COMPLETE - Summary Report V1")
    logger.info(f"⏱️  Total time: {duration}")
    logger.info(f"✅ Success: {len(results) - len(failed_items)} articles")
    logger.info(f"❌ Failed: {len(failed_items)} articles")
    
    if failed_items:
        logger.info("-" * 30)
        logger.info("LIST OF FAILED ITEMS:")
        for idx, f in enumerate(failed_items, 1):
            logger.info(f" {idx}. {f['headline'][:50]}... (Reason: {f['reason']})")
    logger.info("=" * 60)


if __name__ == "__main__":
    main()
